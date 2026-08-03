from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"D:\RPAweb")
DOCS = ROOT / "docs"
SHOTS = DOCS / "_manual_work" / "screenshots"
OUTPUT = DOCS / "AI Tools web员工操作手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "202124"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LINE = "D7DEE8"
GREEN = "E7F6EC"
AMBER = "FFF4D6"
VIOLET = "ECEAFC"
GRAY = "F0F0EC"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.0

    for name, size, color, bold in [
        ("Cover Kicker", 10, BLUE, True),
        ("Cover Title", 29, DARK_BLUE, True),
        ("Cover Subtitle", 14, MUTED, False),
        ("Small Note", 9.5, MUTED, False),
        ("Table Text", 9.5, TEXT, False),
    ]:
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15


def add_custom_numbering(doc, num_fmt, level_text, font_name=None):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    lvl.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    lvl.append(justification)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)

    if font_name:
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        fonts.set(qn("w:eastAsia"), font_name)
        rpr.append(fonts)
        lvl.append(rpr)

    abstract.append(lvl)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    set_numbering(p, num_id)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_callout(doc, label, text, fill=LIGHT_GRAY, border=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.2
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    ppr.append(borders)
    lead = p.add_run(f"{label}  ")
    set_run_font(lead, bold=True, color=DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in [
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ]:
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths_dxa[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def format_table(table, header=True, header_fill=LIGHT_BLUE):
    if header and table.rows:
        set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if header and row_index == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Text"
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                if col_index == 0:
                    for run in paragraph.runs:
                        run.bold = True
                if header and row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)
                else:
                    for run in paragraph.runs:
                        set_run_font(run, size=9.5, bold=(col_index == 0))
    table.rows[0]._tr.get_or_add_trPr()


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_geometry(table, widths_dxa)
    format_table(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def set_picture_alt_text(run, description):
    drawings = run._element.findall(".//" + qn("wp:docPr"))
    for doc_pr in drawings:
        doc_pr.set("descr", description)
        doc_pr.set("title", description)


def add_figure(doc, filename, caption, width=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(SHOTS / filename), width=Inches(width))
    set_picture_alt_text(run, caption)
    cap = doc.add_paragraph(caption, style="Caption")
    return cap


def add_page_heading(doc, text):
    doc.add_page_break()
    return doc.add_paragraph(text, style="Heading 1")


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.text = "AI TOOLS WEB  ·  员工操作手册"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("内部使用  ·  2026-07-27")
    set_run_font(left, size=9, color=MUTED)
    p.add_run("\t")
    add_page_number(p)


def add_body_paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    login_number_id = add_custom_numbering(doc, "decimal", "%1.")
    rpa_number_id = add_custom_numbering(doc, "decimal", "%1.")
    asset_number_id = add_custom_numbering(doc, "decimal", "%1.")
    admin_number_id = add_custom_numbering(doc, "decimal", "%1.")
    bullet_id = add_custom_numbering(doc, "bullet", "•", "Calibri")

    core = doc.core_properties
    core.title = "AI Tools web员工操作手册（精简图文版）"
    core.subject = "员工提交RPA需求、查看进度、提交Skill/Python以及管理员进度管理"
    core.author = "上海焱坤网络科技"
    core.keywords = "AI Tools web, RPA, Skill, Python, 员工操作手册"
    core.comments = "V2.0 精简图文版"

    # 封面：editorial_cover 的克制版变体。
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(10)
    logo_run = logo_p.add_run()
    logo_run.add_picture(str(ROOT / "frontend" / "public" / "logo.png"), width=Inches(0.55))
    set_picture_alt_text(logo_run, "AI Tools web公司标识")

    kicker = doc.add_paragraph("员工操作手册  ·  精简图文版", style="Cover Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)

    title = doc.add_paragraph("AI TOOLS WEB", style="Cover Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)

    subtitle = doc.add_paragraph("5分钟上手：提交需求、查看进度、上传 Skill / Python", style="Cover Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_after = Pt(10)
    run = lead.add_run("正式员工、实习生及管理员的日常操作速查")
    set_run_font(run, size=11, color=TEXT, bold=True)

    add_figure(
        doc,
        "01-home.jpg",
        "图 1  公开首页：RPA、Skill、Python 数量与入口一目了然",
        width=6.1,
    )

    meta = doc.add_paragraph("版本 V2.0  |  更新日期 2026-07-27  |  内网入口 http://172.16.50.20:8090/")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.style = "Small Note"
    add_callout(
        doc,
        "先记住",
        "普通员工点击“用户端入口”；管理员点击“管理员入口”。地址如有调整，以公司最新通知为准。",
    )

    # 1 登录
    add_page_heading(doc, "01  登录与快速开始")
    add_body_paragraph(doc, "首次使用先注册；日常使用可以用用户名或邮箱登录。")
    for text in [
        "打开系统首页，点击右上角“用户端入口”。",
        "没有账号时点击“立即注册”，填写用户名、邮箱和至少 6 位密码。",
        "登录时选择“实习生”或“正式员工”，再输入密码。",
        "登录成功后进入用户工作台，选择需要办理的功能。",
    ]:
        add_list_item(doc, text, login_number_id)
    add_callout(
        doc,
        "职位提示",
        "职位会在首次成功登录后固定。选错或需要变更时联系管理员；不要反复注册新账号。",
        fill=AMBER,
        border="C58B00",
    )
    add_figure(doc, "02-login.jpg", "图 2  用户登录页：选择职位后输入账号和密码", width=5.75)
    add_callout(
        doc,
        "安全提示",
        "共享电脑不要勾选“记住账号和密码”；反馈问题时不要发送密码、验证码或客户隐私数据。",
    )

    # 2 工作台
    add_page_heading(doc, "02  认识员工工作台")
    add_body_paragraph(doc, "顶部菜单与中间卡片作用相同，点击任意一个都可以进入对应功能。")
    add_figure(doc, "03-workbench.jpg", "图 3  用户工作台：五个常用入口", width=6.05)
    add_table(
        doc,
        ["入口", "什么时候使用"],
        [
            ["上传RPA程序", "提交新的自动化开发需求。"],
            ["项目进度", "查看需求审核、RPA开发进度、Skill/Python审核和生命周期。"],
            ["维护记录", "查看历史维护工单及资产拒绝理由。"],
            ["上传 Skill", "登记 Skill 文件名、版本和用途，等待审核。"],
            ["上传插件", "登记 Python 插件文件名、版本、依赖和用法。"],
        ],
        [2250, 7110],
    )
    add_callout(
        doc,
        "实习生",
        "具备当周学习周报权限时会看到“RPA学习情况记录”；没有入口时联系 HR 或管理员核对名单。",
    )

    # 3 RPA需求
    add_page_heading(doc, "03  提交 RPA 开发需求")
    add_body_paragraph(doc, "把流程、输入、输出和验收标准一次写清，可以减少往返确认。")
    add_figure(doc, "04-rpa-request.jpg", "图 4  上传 RPA 程序：带 * 的字段必须填写", width=6.05)
    for text in [
        "填写完整部门、真实姓名、反馈时间和紧急程度。",
        "需求标题按“部门-需求简介-日报/周报/月报”填写。",
        "需求描述写清当前流程、输入数据、操作步骤、期望输出、频率和验收标准。",
        "检查期望完成时间后提交；看到“需求提交成功”才算完成。",
    ]:
        add_list_item(doc, text, rpa_number_id)
    add_callout(
        doc,
        "资料放哪里",
        r"录屏、说明和样例保存到 Z:\运营共享\RPAweb项目\部门，并用需求标题创建文件夹。页面如有新提示，以页面提示为准。",
        fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "账号信息",
        "优先提供低权限测试账号；不要提供个人主账号或与其他系统共用的密码。",
        fill=AMBER,
        border="C58B00",
    )

    # 4 Skill/Python
    add_page_heading(doc, "04  上传 Skill 或 Python 插件")
    add_body_paragraph(doc, "两类资产的操作基本相同：先保存文件，再在系统中登记完整文件名。")
    add_figure(doc, "05-skill-upload.jpg", "图 5  Skill 登记页：当前只登记文件名，不上传文件内容", width=6.05)
    for text in [
        "先把文件保存到管理员通知的约定位置。",
        "填写名称、部门、提交人、版本、说明和包含扩展名的完整文件名。",
        "提交后在当前页面或“项目进度”查看审核结果。",
    ]:
        add_list_item(doc, text, asset_number_id)
    add_table(
        doc,
        ["类型", "命名与说明", "支持格式"],
        [
            ["Skill", "名称建议“部门-功能-skill”；写明用途、触发方式、使用前提。", ".md / .txt / .json / .yaml / .yml / .zip"],
            ["Python", "名称建议“部门-功能-插件”；写明Python版本、依赖、入口和用法。", ".py / .zip / .whl"],
        ],
        [1500, 4760, 3100],
    )
    add_callout(
        doc,
        "审核结果",
        "待审核：等待管理员处理；已通过：进入开发进度和公开看板；已拒绝：查看理由，修正后重新提交。",
    )

    # 5 个人进度
    add_page_heading(doc, "05  查看我的项目进度")
    add_body_paragraph(doc, "进入“项目进度”，一页查看我的需求、RPA项目、维护任务、Skill和Python插件。")
    add_figure(doc, "06-my-progress.jpg", "图 6  项目进度页：顶部是数量，下面是分类明细", width=6.05)
    for text in [
        "“我的需求”看待审核、已通过或已拒绝。",
        "“进行中的项目/已完成的项目”看进度百分比和开发日志。",
        "“我的 Skill 文件/我的 Python 插件”同时显示审核状态、开发进度和生命周期。",
        "数据看不到时先刷新，并确认当前登录账号正确。",
    ]:
        add_list_item(doc, text, bullet_id)
    add_callout(
        doc,
        "两个状态不要混淆",
        "审核状态回答“能不能进入开发”；生命周期回答“项目现在处于什么使用阶段”。",
        fill=LIGHT_BLUE,
    )

    # 6 状态
    add_page_heading(doc, "06  状态速查")
    doc.add_paragraph("审核状态", style="Heading 2")
    add_table(
        doc,
        ["状态", "简单理解", "你要做什么"],
        [
            ["待审核", "已提交，管理员还未处理", "等待审核，不要重复提交"],
            ["已通过", "内容符合要求", "到项目进度查看后续"],
            ["已拒绝", "信息或文件需要修改", "查看理由，修正后重新提交"],
        ],
        [1700, 3560, 4100],
    )
    doc.add_paragraph("开发生命周期", style="Heading 2")
    add_table(
        doc,
        ["状态", "含义", "常见情况"],
        [
            ["在编", "正在开发或尚未达到完成状态", "进度通常为 0%-99%"],
            ["使用", "已投入使用", "自动模式下进度达到 100%"],
            ["大修", "需要较大范围改造", "由管理员手动指定"],
            ["停用", "当前不再使用", "由管理员手动指定，需要恢复时联系负责人"],
        ],
        [1700, 3560, 4100],
    )
    add_callout(
        doc,
        "首页统计",
        "公开首页的项目总数 = RPA + Skill + Python；Skill/Python 通过审核后会显示进度和生命周期。",
        fill=GREEN,
        border="2E8B57",
    )
    doc.add_paragraph("维护记录与学习周报", style="Heading 2")
    add_list_item(doc, "维护记录：点击“详情”查看维护日期、负责人和处理说明。", bullet_id)
    add_list_item(doc, "实习生学习周报：先保存草稿，确认无误后正式提交；被退回时按原因修改并重新提交。", bullet_id)

    # 7 管理员附录
    add_page_heading(doc, "07  管理员快速操作")
    add_callout(
        doc,
        "适用对象",
        "本页仅供有管理员权限的人员使用。普通员工无需进入管理员端。",
    )
    add_figure(
        doc,
        "07-admin-progress.jpg",
        "图 7  开发进度管理：五张统计卡、三类项目、搜索与进度更新",
        width=6.05,
    )
    for text in [
        "需求审核：在 RPA、Skill、Python 标签间切换，查看详情后选择通过或拒绝。",
        "开发进度：五张卡片分别统计项目总数、在编、使用、大修、停用，并显示三类项目明细。",
        "搜索：可按项目名称、部门、状态、提交人或版本查找。",
        "更新：选择进度；状态可随进度自动，也可手动设为在编、使用、大修或停用。",
    ]:
        add_list_item(doc, text, admin_number_id)
    add_callout(
        doc,
        "保存前确认",
        "手动状态会直接影响首页和统计口径。选择“大修”或“停用”前，应确认负责人和原因。",
        fill=AMBER,
        border="C58B00",
    )

    # 8 FAQ/checklist
    add_page_heading(doc, "08  常见问题与提交检查")
    doc.add_paragraph("常见问题", style="Heading 2")
    add_table(
        doc,
        ["问题", "处理方法"],
        [
            ["无法登录", "核对用户名/邮箱、密码和职位；职位已固定时选择正确职位。"],
            ["提交后首页没有显示", "待审核或已拒绝内容不会进入项目统计；先到“项目进度”看审核状态。"],
            ["文件无法上传", "Skill/Python 当前只登记文件名；RPA资料按页面提示保存到公盘。"],
            ["页面没有新数据", "刷新页面，确认账号和公司网络；仍异常时记录页面、时间和错误提示。"],
            ["资产被拒绝", "在项目进度或维护记录查看理由，修改后重新提交。"],
        ],
        [2600, 6760],
    )
    doc.add_paragraph("提交前检查", style="Heading 2")
    checks = [
        "账号和职位正确。",
        "部门、姓名、名称、时间和说明填写完整。",
        "RPA资料已放到正确公盘文件夹。",
        "Skill/Python文件名包含正确扩展名。",
        "没有在标题、截图或附件中暴露密码和不必要的隐私信息。",
        "提交后已在“项目进度”确认状态。",
    ]
    for item in checks:
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        "反馈问题时提供",
        "用户名（不要提供密码）、所属部门、页面名称、操作时间、完整错误提示和脱敏截图。",
        fill=LIGHT_BLUE,
    )
    ending = doc.add_paragraph("本手册依据 2026-07-27 当前系统界面编制。页面字段、权限或路径如有调整，以系统最新提示和公司通知为准。")
    ending.style = "Small Note"
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 避免表格跨页时标题与第一行脱节。
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
